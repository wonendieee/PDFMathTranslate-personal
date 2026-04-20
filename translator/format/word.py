"""Word document format handler - full pipeline (docx -> docx)."""

from __future__ import annotations

import logging
import time
from collections.abc import AsyncGenerator
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING

from translator.format.base import DocumentFormat
from translator.format.base import FormatHandler
from translator.format.word_pipeline.collector import collect_translation_units
from translator.format.word_pipeline.writer import apply_translations
from translator.office.batch_translator import OfficeBatchTranslator

if TYPE_CHECKING:
    from translator.config.model import SettingsModel

logger = logging.getLogger(__name__)


def get_translator(settings):
    """Lazy wrapper around translator.engines.get_translator to avoid
    circular imports (engines -> config.model -> ... -> format.word).

    Tests can patch ``translator.format.word.get_translator`` directly.
    """
    from translator.engines import get_translator as _g

    return _g(settings)


@dataclass
class WordTranslateResult:
    """Mimics babeldoc TranslationResult shape for event consumers."""

    original_path: str
    translated_path: str
    total_seconds: float
    mono_pdf_path: str | None = None
    dual_pdf_path: str | None = None

    @property
    def original_pdf_path(self) -> str:
        return self.original_path


class WordFormatHandler(FormatHandler):
    """.docx handler: paragraph-level translation with dominant-run writeback."""

    def get_format(self) -> DocumentFormat:
        return DocumentFormat.DOCX

    def detect_format(self, file_path: Path) -> bool:
        if not file_path.exists():
            return False
        if file_path.suffix.lower() != ".docx":
            return False
        try:
            import zipfile

            with zipfile.ZipFile(file_path, "r") as zf:
                return any(n.startswith("word/") for n in zf.namelist())
        except Exception:
            return False

    def validate_file(self, file_path: Path) -> bool:
        if not self.detect_format(file_path):
            return False
        try:
            import docx

            doc = docx.Document(file_path)
            _ = doc.paragraphs
            return True
        except Exception as e:
            logger.debug(f".docx validation failed for {file_path}: {e}")
            return False

    async def translate(
        self,
        input_file: Path,
        settings: SettingsModel,
    ) -> AsyncGenerator[dict, None]:
        import docx

        start = time.perf_counter()

        yield {"type": "progress", "overall_progress": 0.05, "stage": "reading"}
        doc = docx.Document(input_file)

        yield {"type": "progress", "overall_progress": 0.10, "stage": "collecting"}
        units = collect_translation_units(doc)
        texts = [u.text for u in units]
        logger.info(f"Collected {len(texts)} translation units from {input_file}")

        translator_obj = get_translator(settings)
        qps = settings.translation.qps
        pool_workers = settings.translation.pool_max_workers or qps
        bt = OfficeBatchTranslator(
            translator_obj,
            qps=qps,
            max_workers=pool_workers,
        )

        if texts:
            translations = await bt.translate_batch(texts)
        else:
            translations = []

        yield {"type": "progress", "overall_progress": 0.85, "stage": "translating"}

        yield {"type": "progress", "overall_progress": 0.95, "stage": "writing"}
        apply_translations(units, translations)

        output_dir = (
            Path(settings.translation.output)
            if settings.translation.output
            else input_file.parent
        )
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{input_file.stem}.zh.docx"
        doc.save(output_path)

        elapsed = time.perf_counter() - start
        result = WordTranslateResult(
            original_path=str(input_file),
            translated_path=str(output_path),
            total_seconds=elapsed,
        )

        yield {
            "type": "finish",
            "translate_result": result,
            "token_usage": {},
        }
