"""Collect translation units from a python-docx Document.

A translation unit = one paragraph's concatenated run text.
Tables are walked: each cell's paragraphs are also units.
Paragraphs containing OMML math are skipped (spec stage 1).
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from typing import TYPE_CHECKING

from translator.office.text_utils import should_translate

if TYPE_CHECKING:
    from docx.document import Document
    from docx.text.paragraph import Paragraph

logger = logging.getLogger(__name__)

_OMML_NS_URIS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/math",
)
_OMML_LOCAL_NAMES = ("oMath", "oMathPara")


@dataclass
class WordTextUnit:
    """A paragraph-level translation unit.

    paragraph_element: the underlying lxml <w:p> element (for injection back)
    run_count: total number of runs at collection time (for diagnostics)
    dominant_run_index: index of the run chosen to receive translated text
    text: concatenated run text (what we send to the translator)
    """

    paragraph_element: object
    run_count: int
    dominant_run_index: int
    text: str


def _has_omml(paragraph: Paragraph) -> bool:
    """Return True if the paragraph XML contains any OMML math element."""
    p = paragraph._p
    for child in p.iter():
        tag = child.tag
        if not isinstance(tag, str) or "}" not in tag:
            continue
        uri = tag.split("}", 1)[0].lstrip("{")
        local = tag.split("}", 1)[1]
        if uri in _OMML_NS_URIS and local in _OMML_LOCAL_NAMES:
            return True
    return False


def _paragraph_to_unit(paragraph: Paragraph) -> WordTextUnit | None:
    """Build a translation unit for one paragraph, or None if it should be skipped."""
    if _has_omml(paragraph):
        logger.debug("Skipping paragraph with OMML math")
        return None

    runs = paragraph.runs
    if not runs:
        return None

    concatenated = "".join(r.text for r in runs)
    if not should_translate(concatenated):
        return None

    dominant_idx = 0
    max_len = len(runs[0].text)
    for i, r in enumerate(runs[1:], start=1):
        if len(r.text) > max_len:
            max_len = len(r.text)
            dominant_idx = i

    return WordTextUnit(
        paragraph_element=paragraph._p,
        run_count=len(runs),
        dominant_run_index=dominant_idx,
        text=concatenated,
    )


def collect_translation_units(doc: Document) -> list[WordTextUnit]:
    """Walk the document and return all translation units in visiting order.

    Order: body paragraphs first (top-to-bottom), then table cells (row-major).
    """
    units: list[WordTextUnit] = []

    for p in doc.paragraphs:
        u = _paragraph_to_unit(p)
        if u is not None:
            units.append(u)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    u = _paragraph_to_unit(p)
                    if u is not None:
                        units.append(u)

    return units
