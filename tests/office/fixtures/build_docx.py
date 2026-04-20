"""Generate Word fixture files for tests.

Run once: `python tests/office/fixtures/build_docx.py`
Output files are committed to the repo (small binary <20KB each).
"""

from pathlib import Path

from docx import Document
from lxml import etree

HERE = Path(__file__).parent


def build_plain():
    doc = Document()
    doc.add_paragraph("Hello, world.")
    doc.add_paragraph("This is a second paragraph.")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Third real paragraph with numbers 123.")
    doc.save(HERE / "plain.docx")


def build_with_table():
    doc = Document()
    doc.add_paragraph("Title paragraph.")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Row 1 Col 1"
    t.cell(0, 1).text = "Row 1 Col 2"
    t.cell(1, 0).text = "Row 2 Col 1"
    t.cell(1, 1).text = "Row 2 Col 2"
    doc.add_paragraph("Footer paragraph.")
    doc.save(HERE / "with_table.docx")


def build_mixed_runs():
    """Paragraph with bold + normal + italic runs."""
    doc = Document()
    p = doc.add_paragraph("")
    p.add_run("Normal ")
    p.add_run("bold").bold = True
    p.add_run(" and ")
    p.add_run("italic").italic = True
    p.add_run(" here.")
    doc.save(HERE / "mixed_runs.docx")


def build_with_formula():
    """Paragraph that is purely an OMML math formula."""
    doc = Document()
    doc.add_paragraph("Before formula.")
    p = doc.add_paragraph()
    omath_xml = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '<m:r><m:t>x^2 + y^2 = z^2</m:t></m:r></m:oMath>'
    )
    p._p.append(etree.fromstring(omath_xml))
    doc.add_paragraph("After formula.")
    doc.save(HERE / "with_formula.docx")


if __name__ == "__main__":
    build_plain()
    build_with_table()
    build_mixed_runs()
    build_with_formula()
    print("Fixtures generated in:", HERE)
