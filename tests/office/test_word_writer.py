"""Tests for Word dominant-run writeback."""
from pathlib import Path

from docx import Document
from translator.format.word_pipeline.collector import collect_translation_units
from translator.format.word_pipeline.writer import apply_translations

FIXTURES = Path(__file__).parent / "fixtures"


def test_writeback_plain_replaces_paragraph_text(tmp_path):
    src = FIXTURES / "plain.docx"
    doc = Document(src)
    units = collect_translation_units(doc)
    translations = [u.text.upper() for u in units]
    apply_translations(units, translations)

    out = tmp_path / "out.docx"
    doc.save(out)

    d2 = Document(out)
    all_text = "\n".join(p.text for p in d2.paragraphs)
    assert "HELLO, WORLD." in all_text
    assert "THIS IS A SECOND PARAGRAPH." in all_text


def test_writeback_mixed_runs_preserves_run_count(tmp_path):
    doc = Document(FIXTURES / "mixed_runs.docx")
    units = collect_translation_units(doc)
    assert len(units) == 1
    original_run_count = units[0].run_count
    apply_translations(units, ["TRANSLATED"])

    out = tmp_path / "mixed_out.docx"
    doc.save(out)
    d2 = Document(out)
    p = d2.paragraphs[0]
    assert len(p.runs) == original_run_count
    assert p.runs[0].text == "TRANSLATED"
    for r in p.runs[1:]:
        assert r.text == ""


def test_writeback_preserves_dominant_run_formatting(tmp_path):
    doc = Document(FIXTURES / "mixed_runs.docx")
    units = collect_translation_units(doc)
    apply_translations(units, ["TRANS"])

    out = tmp_path / "mixed2.docx"
    doc.save(out)
    d2 = Document(out)
    assert d2.paragraphs[0].runs[0].bold is not True
