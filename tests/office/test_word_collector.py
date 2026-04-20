"""Tests for Word paragraph text collection."""
from pathlib import Path

from docx import Document
from translator.format.word_pipeline.collector import WordTextUnit
from translator.format.word_pipeline.collector import collect_translation_units

FIXTURES = Path(__file__).parent / "fixtures"


def test_collect_plain_yields_non_empty_paragraphs_only():
    doc = Document(FIXTURES / "plain.docx")
    units = collect_translation_units(doc)
    texts = [u.text for u in units]
    assert "Hello, world." in texts
    assert "This is a second paragraph." in texts
    assert "Third real paragraph with numbers 123." in texts
    assert "" not in texts
    assert "   " not in texts


def test_collect_table_cells():
    doc = Document(FIXTURES / "with_table.docx")
    units = collect_translation_units(doc)
    texts = [u.text for u in units]
    assert "Title paragraph." in texts
    assert "Row 1 Col 1" in texts
    assert "Row 2 Col 2" in texts
    assert "Footer paragraph." in texts


def test_collect_skips_pure_formula_paragraph():
    doc = Document(FIXTURES / "with_formula.docx")
    units = collect_translation_units(doc)
    texts = [u.text for u in units]
    assert "Before formula." in texts
    assert "After formula." in texts
    assert not any("x^2" in t for t in texts)


def test_collect_mixed_runs_produces_single_unit_with_concatenated_text():
    doc = Document(FIXTURES / "mixed_runs.docx")
    units = collect_translation_units(doc)
    assert len(units) == 1
    u = units[0]
    assert u.text == "Normal bold and italic here."
    assert u.dominant_run_index == 0


def test_word_text_unit_has_required_fields():
    u = WordTextUnit(
        paragraph_element=None,
        run_count=3,
        dominant_run_index=1,
        text="hello",
    )
    assert u.text == "hello"
    assert u.run_count == 3
    assert u.dominant_run_index == 1
