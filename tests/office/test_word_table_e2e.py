"""End-to-end: table cells survive the collector+writer pipeline."""
from pathlib import Path

from docx import Document
from translator.format.word_pipeline.collector import collect_translation_units
from translator.format.word_pipeline.writer import apply_translations

FIXTURES = Path(__file__).parent / "fixtures"


def test_table_cells_translated(tmp_path):
    doc = Document(FIXTURES / "with_table.docx")
    units = collect_translation_units(doc)
    assert len(units) == 6

    translations = [f"[T]{u.text}" for u in units]
    apply_translations(units, translations)

    out = tmp_path / "table_out.docx"
    doc.save(out)

    d2 = Document(out)
    assert d2.paragraphs[0].text.startswith("[T]")
    cells_text = []
    for table in d2.tables:
        for row in table.rows:
            for cell in row.cells:
                cells_text.append(cell.text)
    assert all(t.startswith("[T]") for t in cells_text)
    assert len(cells_text) == 4


def test_table_structure_preserved(tmp_path):
    doc = Document(FIXTURES / "with_table.docx")
    units = collect_translation_units(doc)
    apply_translations(units, ["X"] * len(units))

    out = tmp_path / "struct.docx"
    doc.save(out)

    d2 = Document(out)
    assert len(d2.tables) == 1
    t = d2.tables[0]
    assert len(t.rows) == 2
    assert len(t.columns) == 2
