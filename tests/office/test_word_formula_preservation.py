"""OMML paragraphs must pass through unchanged (spec stage 1 guarantee)."""
from pathlib import Path

from docx import Document
from lxml import etree
from translator.format.word_pipeline.collector import collect_translation_units
from translator.format.word_pipeline.writer import apply_translations

FIXTURES = Path(__file__).parent / "fixtures"


def test_formula_paragraph_not_in_units():
    doc = Document(FIXTURES / "with_formula.docx")
    units = collect_translation_units(doc)
    assert all("x^2" not in u.text for u in units)


def test_formula_paragraph_text_unchanged_after_pipeline(tmp_path):
    src = FIXTURES / "with_formula.docx"
    doc = Document(src)

    formula_paragraphs_before = []
    for p in doc.paragraphs:
        has_omml = any(
            (isinstance(c.tag, str) and c.tag.endswith("}oMath"))
            for c in p._p.iter()
        )
        if has_omml:
            formula_paragraphs_before.append(etree.tostring(p._p))

    assert len(formula_paragraphs_before) == 1

    units = collect_translation_units(doc)
    apply_translations(units, [u.text.upper() for u in units])

    out = tmp_path / "f.docx"
    doc.save(out)
    d2 = Document(out)

    formula_paragraphs_after = []
    for p in d2.paragraphs:
        has_omml = any(
            (isinstance(c.tag, str) and c.tag.endswith("}oMath"))
            for c in p._p.iter()
        )
        if has_omml:
            formula_paragraphs_after.append(etree.tostring(p._p))

    assert formula_paragraphs_after == formula_paragraphs_before
